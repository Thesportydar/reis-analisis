import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from wordcloud import WordCloud
import pandas as pd
import numpy as np
import re
import distinctipy

import nltk
from nltk.corpus import stopwords
import os
from datetime import datetime
import sqlite3
import uuid

from Constants import FILE_PREFIX, REPORT_PREFIX
from dotenv import load_dotenv

load_dotenv()
UPLOADS_DIR = os.getenv("UPLOADS_DIR", "/var/lib/reis")
DB_PATH = os.getenv("DATABASE_URL", "sqlite.db")


class ReportGenerator:
    def __init__(self, rei_id, period, df_info_reis_pregs, df_info_reis_pregnuevas, num_clusters):
        self.rei_id = rei_id
        self.period = period
        self.df_info_reis_pregs = df_info_reis_pregs
        self.corpus = df_info_reis_pregs['pregunta'].tolist()
        self.df_info_reis_pregnuevas = df_info_reis_pregnuevas
        self.num_clusters = num_clusters if num_clusters else 0
        self.document = docx.Document()
        self.set_styles()
        self.colors = distinctipy.get_colors(self.num_clusters)
        self.init_clusters()
        self.init_nltk()


    def set_styles(self):
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)


    def init_clusters(self):
        self.df_info_reis_clusters = pd.DataFrame()
        self.df_info_reis_clusters['color']= self.colors
        self.df_info_reis_clusters = self.df_info_reis_clusters.reset_index(drop=True)
        self.df_info_reis_clusters.index = self.df_info_reis_clusters.index + 1


    def init_nltk(self):
        nltk.download('punkt')
        nltk.download('punkt_tab')
        nltk.download('stopwords')
        self.stopwords = set(stopwords.words('spanish', 'english'))
        self.stopwords.update([ "pregunta", "preguntas", "clase", "cúal", "cúales", "cómo"])


    @staticmethod
    def RGBtoHex(vals, rgbtype=1):
        """Converts RGB values in a variety of formats to Hex values.

        @param  vals     An RGB/RGBA tuple
        @param  rgbtype  Valid valus are:
                1 - Inputs are in the range 0 to 1
                256 - Inputs are in the range 0 to 255

        @return A hex string in the form '#RRGGBB' or '#RRGGBBAA'
        """

        if len(vals)!=3 and len(vals)!=4:
            raise Exception("RGB or RGBA inputs to RGBtoHex must have three or four elements!")
        if rgbtype!=1 and rgbtype!=256:
            raise Exception("rgbtype must be 1 or 256!")

        #Convert from 0-1 RGB/RGBA to 0-255 RGB/RGBA
        if rgbtype==1:
            vals = [255*x for x in vals]

        #Ensure values are rounded integers, convert to hex, and concatenate
        return '#' + ''.join(['{:02X}'.format(int(round(x))) for x in vals])


    def add_heading(self, text, level=2):
        self.document.add_heading(text, level)


    def add_paragraph(self, text):
        self.document.add_paragraph(text)


    def add_table(self, rows, cols, column_widths):
        table = self.document.add_table(rows=rows, cols=cols)
        for i, width in enumerate(column_widths):
            table.columns[i].width = Inches(width)
        return table


    @staticmethod
    def process_string(input_string):
        return re.sub(r",", "\n", input_string)

    def generate_wordcloud(self, text, output_file):
        wordcloud = WordCloud(background_color="white").generate(text)
        wordcloud.to_file(output_file)
        self.document.add_picture(output_file, width=Inches(2.5))


    def add_questions_section(self):
        self.add_heading('Preguntas', 2)

        all_q_text = ' \n'
        for idx, sentence in enumerate(self.corpus):
            all_q_text = all_q_text + sentence + '\n'

        self.add_paragraph(all_q_text)


    def add_clusters_section(self):
        self.document.add_page_break()
        self.add_heading('Preguntas por temas')
        self.add_paragraph(' ')

        table = self.add_table(rows = 1, cols = 3, column_widths = [0.6, 0.6, 5.8])

        #adding headers rows
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Tema'
        hdr_cells[1].text = 'Votos'
        hdr_cells[2].text = 'Pregunta'

        for i in range(1, self.num_clusters+1):
            df_c = self.df_info_reis_pregs.loc[self.df_info_reis_pregs['cluster'] == i]

            joined_string = ' '.join(df_c['pregunta'].astype(str))

            for index, row in df_c.iterrows():
                #second riw
                row_cells = table.add_row().cells
                row_cells[0].text = ""
                rgb = self.df_info_reis_clusters.loc[i, 'color']
                background_color = ReportGenerator.RGBtoHex(rgb,1)

                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
                row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)

                row_cells[1].text = str(row['votos'])
                row_cells[2].text = row['pregunta']


    def add_wordcloud_clusters_section(self):
        self.add_heading('Preguntas por temas')

        for i in range(1, self.num_clusters+1):
            df_c = self.df_info_reis_pregs.loc[self.df_info_reis_pregs['cluster'] == i]

            # si un cluster no tiene preguntas, se salta y continua con el siguiente
            if df_c.empty:
                continue

            rgb = self.df_info_reis_clusters.loc[i, 'color']
            background_color = ReportGenerator.RGBtoHex(rgb,1)

            #heading
            table = self.add_table(rows = 1, cols = 1, column_widths = [7])
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Tema '+str(i)
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
            hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm)

            # Generate a word cloud image
            joined_string = ' '.join(df_c['pregunta'].astype(str))

            # WARNING: si todas las palabras son stopwords lanza un error
            wordcloud = WordCloud(stopwords=self.stopwords, background_color="white").generate(joined_string)
            wordcloud.to_file("image.png") # Save the image to a file

            # Add the word cloud image to the document
            self.document.add_picture('image.png', width=Inches(2.5))

            table = self.add_table(rows = 1, cols = 3, column_widths = [0.6, 0.6, 5.8])

            #adding headers rows
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'Grupos'
            hdr_cells[2].text = 'Pregunta'
            shading_elm = parse_xml(r'<w:shd {} w:fill="dce3db"/>'.format(nsdecls('w')))
            hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
            shading_elm = parse_xml(r'<w:shd {} w:fill="dce3db"/>'.format(nsdecls('w')))
            hdr_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
            shading_elm = parse_xml(r'<w:shd {} w:fill="dce3db"/>'.format(nsdecls('w')))
            hdr_cells[2]._tc.get_or_add_tcPr().append(shading_elm)

            for index, row in df_c.iterrows():
                #second row
                row_cells = table.add_row().cells

                row_cells[0].text = str(row['votos'])
                text_grupos = ReportGenerator.process_string(str(row['grupo']))
                row_cells[1].text = text_grupos
                row_cells[2].text = row['pregunta']

            self.document.add_paragraph("___________________________________________________________________")


    def add_page_cluster_section(self):
        all_text = ""

        for i in range(1, self.num_clusters+1):
            # si un cluster no tiene preguntas, se salta y continua con el siguiente
            if self.df_info_reis_pregs.loc[self.df_info_reis_pregs['cluster'] == i].empty:
                continue

            self.document.add_page_break()

            rgb = self.df_info_reis_clusters.loc[i, 'color']
            background_color = ReportGenerator.RGBtoHex(rgb,1)

            #heading
            table = self.add_table(rows = 1, cols = 1, column_widths = [7])
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Tema '+str(i)

            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
            hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm)

            #preguntas
            self.add_heading('Preguntas', 2)
            #  document.add_paragraph(' ')

            table = self.add_table(rows = 1, cols = 2, column_widths = [0.6, 6.4])

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Votos'
            hdr_cells[1].text = 'Pregunta'
            shading_elm = parse_xml(r'<w:shd {} w:fill="dce3db"/>'.format(nsdecls('w')))
            hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
            shading_elm = parse_xml(r'<w:shd {} w:fill="dce3db"/>'.format(nsdecls('w')))
            hdr_cells[1]._tc.get_or_add_tcPr().append(shading_elm)

            df_c = self.df_info_reis_pregs.loc[self.df_info_reis_pregs['cluster'] == i]

            for index, row in df_c.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = ""
                rgb = self.df_info_reis_clusters.loc[i, 'color']
                background_color = ReportGenerator.RGBtoHex(rgb,1)
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
                row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
                row_cells[0].text = str(row['votos'])
                row_cells[1].text = row['pregunta']

            #nuevas preguntas
            self.add_heading('Nuevas preguntas', 2)
            df_c = self.df_info_reis_pregnuevas.loc[self.df_info_reis_pregnuevas['cluster'] == i]

            if df_c['pregunta'].str.strip().eq('').all():
                self.add_paragraph("No hay nuevas preguntas.")
            else:
                table = self.add_table(rows = 1, cols = 2, column_widths = [1.2, 5.8])

                hdr_cells = table.rows[0].cells
                shading_elm = parse_xml(r'<w:shd {} w:fill="dce3db"/>'.format(nsdecls('w')))
                hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
                shading_elm = parse_xml(r'<w:shd {} w:fill="dce3db"/>'.format(nsdecls('w')))
                hdr_cells[1]._tc.get_or_add_tcPr().append(shading_elm)
                hdr_cells[0].text = 'Grupo'
                hdr_cells[1].text = 'Nuevas preguntas'


                for index, row in df_c.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = ""
                    rgb = self.df_info_reis_clusters.loc[i, 'color']
                    background_color = ReportGenerator.RGBtoHex(rgb,1)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
                    row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
                    row_cells[0].text = str(row['grupo'])
                    row_cells[1].text = row['pregunta']


    def save_report(self):
        # Nombre con el cual el usuario conoce el archivo
        original_name = f"{self.rei_id}_{self.period}_reporte.docx"
        # Nombre con el cual se guarda el archivo
        save_name = REPORT_PREFIX + "-" + datetime.now().strftime("%Y%m%d%H%M%S") + ".docx"
        # Ruta donde se guarda el archivo
        save_path = os.path.join(UPLOADS_DIR, save_name)
        # URL para descargar el archivo
        download_url = os.path.join(FILE_PREFIX, save_name)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # Asegurar que la carpeta de uploads exista
        os.makedirs(UPLOADS_DIR, exist_ok=True)

        self.document.save(save_path)
        print(f"Reporte guardado como {save_name}")

        try:
            score = int(self.num_clusters)
        except (ValueError, TypeError):
            score = 0


        # guarda los metadatos en la base de datos
        query = """
            INSERT INTO report (id, class_id, fileName, filePath, mime, score)
            VALUES (?, ?, ?, ?, ?, ?)
        """

        report_id = str(uuid.uuid4())

        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.cursor()
            cursor.execute(query, (report_id, self.period, original_name, download_url, mime, score))
            conn.commit()

        return {'file_name': original_name, 'file_path': save_path, 'download_url': download_url, 'mime': mime}


    def generate_report(self):
        self.add_questions_section()
        self.add_clusters_section()
        self.add_wordcloud_clusters_section()
        self.add_page_cluster_section()
        return self.save_report()


if __name__ == '__main__':
    generator = ReportGenerator(
        rei_id="1234",
        period="2025-01",
        df_info_reis_pregs=pd.DataFrame(),
        df_info_reis_pregnuevas=pd.DataFrame(),
        num_clusters=5
    )
    generator.generate_report()
